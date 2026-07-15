"""Resume parsing utilities for reading and parsing PDF/DOCX/TXT files with Hebrew RTL support."""

from __future__ import annotations

import hashlib
import io
import mimetypes
import os
from pathlib import Path

import fitz  # PyMuPDF
import logging
import subprocess
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

def parse_pdf_content(file_content: bytes) -> str:
    """
    Parses PDF content using PyMuPDF (fitz).
    Tries layout-preserving 'blocks' mode first.
    If that produces fragmented text (one char per line), falls back to 'text' mode.
    """
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        full_text = []
        
        # First pass: Try blocks (layout preserving)
        for page in doc:
            # get_text("blocks") returns a list of tuples: (x0, y0, x1, y1, "text", block_no, block_type)
            blocks = page.get_text("blocks")
            
            # Filter for text blocks (type 0) and remove empty ones
            text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
            
            # Sort blocks using layout analysis (Columns vs Rows)
            sorted_blocks = _sort_blocks_by_layout(text_blocks)

            for b in sorted_blocks:
                text_content = b[4].strip()
                if text_content:
                    full_text.append(text_content)

        text_result = "\n\n".join(full_text)
        
        # Check if broken
        if _is_extraction_broken(text_result):
            logger.info("PyMuPDF 'blocks' mode produced fragmented text. Retrying with 'text' mode...")
            full_text = []
            for page in doc:
                # "text" mode: extracts text in natural reading order, handling some layout issues
                # sort=True attempts to sort by vertical position then horizontal
                full_text.append(page.get_text("text", sort=True))
            text_result = "\n".join(full_text)
            
        return text_result

    except Exception as e:
        logger.error(f"Error parsing PDF with PyMuPDF: {e}")
        # Fallback or re-raise depending on your strategy
        raise e


def _sort_blocks_by_layout(blocks: list) -> list:
    """
    Sort blocks by analyzing the page layout (columns vs rows).
    1. Detects if there is a vertical column separator.
    2. If found, splits page into bands (separated by full-width blocks like headers).
    3. Within each band, reads Left Column then Right Column.
    """
    if not blocks:
        return []
        
    # 1. Find page width boundaries
    min_x = min(b[0] for b in blocks)
    max_x = max(b[2] for b in blocks)
    width = max_x - min_x
    
    # 2. Search for a column splitter in the middle 50% of the page
    search_start = min_x + width * 0.25
    search_end = min_x + width * 0.75
    
    best_split = -1
    min_intersect_count = len(blocks) + 1
    
    # Scan X axis with step of 5 to find a vertical gap
    scan_x = search_start
    while scan_x < search_end:
        # Count blocks crossing this line
        count = sum(1 for b in blocks if b[0] < scan_x < b[2])
        if count < min_intersect_count:
            min_intersect_count = count
            best_split = scan_x
        scan_x += 5
        
    # Threshold: if too many blocks cross the split, assume single column
    # Allow up to 3 crossing blocks (e.g. Header, Footer, Horizontal Line)
    # OR up to 10% of total blocks
    threshold = max(3, len(blocks) * 0.1)
    
    if min_intersect_count > threshold:
        # Fallback to standard Y-sort (row by row)
        # Group by Y (rounded) then X
        return sorted(blocks, key=lambda b: (round(b[1] / 10) * 10, b[0]))
        
    # 3. Classify blocks
    spanning = []
    left = []
    right = []
    
    for b in blocks:
        if b[0] < best_split < b[2]:
            spanning.append(b)
        elif b[2] <= best_split:
            left.append(b)
        else:
            right.append(b)
            
    # 4. Create Bands based on Spanning blocks
    # Sort spanning blocks by Y
    spanning.sort(key=lambda b: b[1])
    
    final_order = []
    current_y = -1.0
    
    def get_blocks_in_band(block_list, top, bottom):
        # Use vertical center of block to determine band membership
        return [b for b in block_list if top <= (b[1] + b[3]) / 2 < bottom]

    for sp in spanning:
        sp_top = sp[1]
        sp_bottom = sp[3]
        
        # Process band above this spanning block
        l_band = get_blocks_in_band(left, current_y, sp_top)
        r_band = get_blocks_in_band(right, current_y, sp_top)
        
        # Sort columns by Y
        l_band.sort(key=lambda b: (b[1], b[0]))
        r_band.sort(key=lambda b: (b[1], b[0]))
        
        final_order.extend(l_band)
        final_order.extend(r_band)
        
        # Add the spanning block
        final_order.append(sp)
        
        current_y = sp_bottom
        
    # Process final band
    l_band = get_blocks_in_band(left, current_y, 99999)
    r_band = get_blocks_in_band(right, current_y, 99999)
    
    l_band.sort(key=lambda b: (b[1], b[0]))
    r_band.sort(key=lambda b: (b[1], b[0]))
    
    final_order.extend(l_band)
    final_order.extend(r_band)
    
    return final_order


def parse_text_content(file_content: bytes) -> str:
    """Helper for plain text files"""
    return file_content.decode("utf-8", errors="ignore")


# --- File I/O ---

def sha256_of_bytes(data: bytes) -> str:
    """Return SHA256 hash of bytes."""
    h = hashlib.sha256()
    h.update(data)
    return h.hexdigest()


def detect_mime(path: Path) -> str:
    """Guess MIME type from extension."""
    guess, _ = mimetypes.guess_type(str(path))
    return guess or "application/octet-stream"


def read_file_bytes(path: Path) -> bytes:
    """Read file content as bytes."""
    return path.read_bytes()


# --- Parsing ---

def _reconstruct_text_from_words(pdf_page) -> str:
    """
    Fallback: reconstruct a page's text from words when extract_text() is empty.
    Groups words by their top (y) position and sorts by x to form lines.
    """
    words = pdf_page.extract_words() or []
    if not words:
        return ""
    # Group by y position (rounded to avoid micro-variations)
    rows = {}
    for w in words:
        y = int(round(w.get("top", 0)))
        rows.setdefault(y, []).append(w)
    lines = []
    for y in sorted(rows.keys()):
        parts = sorted(rows[y], key=lambda w: w.get("x0", 0))
        line = " ".join(p.get("text", "") for p in parts if p.get("text"))
        if line.strip():
            lines.append(line.strip())
    return "\n".join(lines)


def parse_to_text(file_path: str) -> str:
    """
    Main entry point to extract text based on file extension.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.doc':
        return extract_text_from_doc(file_path)
    elif ext == '.pdf':
        # Try PyMuPDF (fitz) first via parse_pdf_content
        try:
            file_bytes = read_file_bytes(Path(file_path))
            text = parse_pdf_content(file_bytes)
            
            # Check if result is good
            if text and not _is_extraction_broken(text):
                return text
                
            logger.warning(f"PyMuPDF extraction broken or empty for {file_path}, falling back to pdfplumber")
        except Exception as e:
            logger.error(f"PyMuPDF failed for {file_path}: {e}")
            
        # Fallback to pdfplumber with custom settings
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Try with loose tolerance for Hebrew/spaced text
                    # x_tolerance: horizontal distance to merge chars
                    # y_tolerance: vertical distance to merge lines
                    page_text = page.extract_text(x_tolerance=2, y_tolerance=3)
                    if page_text:
                        text += page_text + "\n"
            
            if _is_extraction_broken(text):
                 logger.info("Standard pdfplumber extraction broken, retrying with high tolerance...")
                 text = ""
                 with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        # Very loose tolerance to force grouping
                        # x_tolerance=15 is very aggressive for merging horizontal chars
                        page_text = page.extract_text(x_tolerance=15, y_tolerance=10)
                        if page_text:
                            text += page_text + "\n"

            return text
        except ImportError:
            logger.warning("pdfplumber not installed")
            return ""
        except Exception as e:
            logger.error(f"Error reading PDF with pdfplumber: {e}")
            return ""
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
            
    return ""


def _is_extraction_broken(text: str) -> bool:
    """
    Heuristic to check if text extraction resulted in one-char-per-line garbage.
    """
    if not text:
        return True
    
    lines = text.strip().split('\n')
    if not lines:
        return True
        
    # Count lines with 1 or 2 characters
    short_lines = sum(1 for line in lines if len(line.strip()) <= 2)
    total_lines = len(lines)
    
    # If more than 40% of lines are 1-2 chars, it's likely broken
    if total_lines > 10 and (short_lines / total_lines) > 0.4:
        return True
        
    return False


def _extract_text_from_xml(element):
    """
    Helper function to extract text from XML element recursively.
    This catches text inside Text Boxes (w:txbxContent) which standard python-docx misses.
    """
    text_parts = []
    
    # Iterate over all elements in the XML tree
    for node in element.iter():
        # Check for text tag <w:t>
        if node.tag.endswith('}t'):
            if node.text:
                text_parts.append(node.text)
        # Check for breaks and paragraphs to add newlines
        elif node.tag.endswith('}br') or node.tag.endswith('}cr'):
            text_parts.append('\n')
        elif node.tag.endswith('}p'):
            text_parts.append('\n')
        elif node.tag.endswith('}tab'):
            text_parts.append('\t')
            
    return "".join(text_parts).strip()


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file.
    Uses XML parsing for Headers to catch Text Boxes/Shapes.
    Uses standard parsing for Body to preserve structure.
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        # 1. Aggressive Header Extraction
        # We use a set to avoid reading the same header twice (if it repeats on every page)
        processed_headers = set()
        
        for section in doc.sections:
            # Check all possible header types
            headers = [section.header, section.first_page_header, section.even_page_header]
            
            for header in headers:
                # Check if header exists and is not linked to previous (if linked, it's the same as previous section's header)
                # Also check if we already processed this header part
                if header and not header.is_linked_to_previous:
                    # Use part as unique identifier
                    if header.part in processed_headers:
                        continue
                    processed_headers.add(header.part)
                    
                    # The magic: direct XML reading from the header
                    # This catches Text Boxes and Shapes that users can't select with Ctrl+A
                    try:
                        header_xml = header.part.element
                        header_text = _extract_text_from_xml(header_xml)
                        if header_text:
                            full_text.append(header_text)
                            full_text.append("-" * 20) # Visual separator
                    except Exception as e:
                        logger.warning(f"Failed to parse header XML: {e}")

        # 2. Extract Body - Preserving order of paragraphs and tables
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Use XML extraction for paragraphs too! 
                # This catches floating text boxes anchored to the paragraph.
                para_text = _extract_text_from_xml(element)
                if para_text:
                    full_text.append(para_text)
            
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

        return "\n".join(full_text)

    except Exception as e:
        logger.error(f"Error reading DOCX file {file_path}: {e}")
        return ""


def extract_text_from_doc(file_path: str) -> str:
    """
    Extracts text from binary .doc files using catdoc (requires catdoc installed on OS).
    """
    try:
        # Use catdoc which is lightweight and fast
        result = subprocess.run(
            ['catdoc', '-w', file_path], 
            capture_output=True, 
            text=True,
            encoding='utf-8',
            errors='ignore'
        )
        if result.returncode == 0:
            return result.stdout
        else:
            logger.error(f"catdoc failed: {result.stderr}")
            return ""
    except FileNotFoundError:
        logger.warning("catdoc tool not found. Cannot process .doc files.")
        return ""
    except Exception as e:
        logger.error(f"Error reading DOC file {file_path}: {e}")
        return ""


def _parse_docx_advanced(file_content: bytes) -> str:
    """
    DEPRECATED: Kept for reference but unused in favor of extract_text_from_docx.
    Parse DOCX with enhanced Hebrew RTL support.
    Uses python-docx as primary parser with table extraction.
    """
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        doc = Document(io.BytesIO(file_content))
        parts = []
        
        # Process document body in order (paragraphs and tables)
        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('}p'):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para and para.text.strip():
                    parts.append(para.text.strip())
            
            # Check if it's a table
            elif element.tag.endswith('}tbl'):
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                if table:
                    # Extract table with better formatting
                    table_text = _extract_table_text(table)
                    if table_text:
                        parts.append(table_text)
        
        # If no structured extraction worked, fallback to simple paragraphs
        if not parts:
            return _parse_docx_fallback(file_content)
        
        return "\n\n".join(parts)
    
    except Exception as e:
        print(f"Advanced DOCX parsing failed: {e}")
        return _parse_docx_fallback(file_content)


def _extract_table_text(table) -> str:
    """
    Extract text from DOCX table with proper structure preservation.
    Handles multi-column layouts common in resumes.
    """
    lines = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                cells.append(cell_text)
        
        if cells:
            # Join cells with separator
            # If it's a single cell, just add it
            if len(cells) == 1:
                lines.append(cells[0])
            else:
                # Multiple cells - likely structured data
                lines.append(" | ".join(cells))
    
    return "\n".join(lines)


def _parse_docx_fallback(file_content: bytes) -> str:
    """
    Fallback DOCX parser using python-docx.
    Extracts paragraphs and tables.
    """
    from docx import Document
    
    doc = Document(io.BytesIO(file_content))
    parts = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))
    
    return "\n".join(parts).strip()
